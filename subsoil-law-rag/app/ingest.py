"""Разбор PDF/DOCX и нарезка текста на фрагменты с привязкой к статьям/пунктам."""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable, List, Optional, Sequence

# Ссылка на источник — это статья или пункт. Раздел и глава хранятся отдельными
# полями (см. app/sections.py), поэтому в локаторы они не попадают.
STRUCTURAL_PATTERNS = [
    re.compile(r"^\s*(Статья\s+\d+(?:[-–]\d+)?(?:\.\d+)*)", re.IGNORECASE),
    re.compile(r"^\s*(Приложение\s+[\dIVXLC]+)", re.IGNORECASE),
    re.compile(r"^\s*(Пункт\s+\d+(?:\.\d+)*)", re.IGNORECASE),
]

# Служебные абзацы: история изменений и пометки составителя базы. Нормами не
# являются, а в поиске стабильно шумят.
NOISE_PATTERNS = [
    re.compile(r"^\s*Сноска[\.\s]", re.IGNORECASE),
    re.compile(r"^\s*Примечание\s+ИЗПИ", re.IGNORECASE),
    re.compile(r"^\s*(СОДЕРЖАНИЕ|Оглавление)\s*$", re.IGNORECASE),
    re.compile(r"^\s*Для удобства пользования", re.IGNORECASE),
]
# Многоуровневый номер в начале абзаца: «3.1.2)».
MULTILEVEL_POINT = re.compile(r"^\s*(\d+(?:\.\d+){1,3})[\.\)]\s+\S")
# Сплошная нумерация пунктов подзаконного акта: «100. В зависимости от...».
SINGLE_POINT = re.compile(r"^\s*(\d{1,4})\.\s+\S")
ARTICLE_HEAD = re.compile(r"^\s*Статья\s+\d+", re.IGNORECASE)

# «Статья»/«Приложение» всегда открывают новый фрагмент.
BOUNDARY_PREFIXES = ("статья", "приложение")

SUPPORTED_SUFFIXES = {".pdf", ".docx", ".txt", ".md"}

# Стили нумерации: кодекс делится на статьи, правила и приказы — на пункты.
NUMBERING_ARTICLES = "articles"
NUMBERING_POINTS = "points"


class UnsupportedFormat(ValueError):
    """Формат файла не поддерживается."""


@dataclass
class Block:
    """Абзац исходного документа с известной страницей и местом в иерархии."""

    text: str
    page: Optional[int] = None
    section: str = ""
    chapter: str = ""


@dataclass
class Chunk:
    """Готовый к индексации фрагмент."""

    text: str
    locator: str = ""
    page: Optional[int] = None
    index: int = 0
    section: str = ""
    chapter: str = ""
    metadata: dict = field(default_factory=dict)


def detect_numbering(blocks: Sequence[Block]) -> str:
    """Определить стиль нумерации документа по числу заголовков «Статья N»."""
    articles = sum(1 for block in blocks if ARTICLE_HEAD.match(block.text))
    return NUMBERING_ARTICLES if articles >= 5 else NUMBERING_POINTS


def detect_locator(line: str, numbering: str = NUMBERING_ARTICLES) -> Optional[str]:
    """Вернуть номер статьи/пункта, если строка выглядит как заголовок."""
    stripped = line.strip()
    if not stripped or len(stripped) > 250:
        return None
    for pattern in STRUCTURAL_PATTERNS:
        match = pattern.match(stripped)
        if match:
            return match.group(1).strip().rstrip(".")
    match = MULTILEVEL_POINT.match(stripped)
    if match:
        return match.group(1)
    if numbering == NUMBERING_POINTS:
        match = SINGLE_POINT.match(stripped)
        if match:
            return f"пункт {match.group(1)}"
    return None


def is_noise(text: str) -> bool:
    """Служебный абзац, который не нужно индексировать."""
    return any(pattern.match(text) for pattern in NOISE_PATTERNS)


def _normalize(text: str) -> str:
    text = text.replace("­", "").replace("\xa0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_pdf(path: Path) -> List[Block]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    blocks: List[Block] = []
    for page_no, page in enumerate(reader.pages, start=1):
        text = _normalize(page.extract_text() or "")
        if not text:
            continue
        for paragraph in text.split("\n"):
            if paragraph.strip():
                blocks.append(Block(text=paragraph.strip(), page=page_no))
    return blocks


def extract_docx(path: Path) -> List[Block]:
    """Абзацы и таблицы в порядке документа.

    ``document.paragraphs`` и ``document.tables`` — два независимых списка, и
    если читать их подряд, все таблицы окажутся в конце, оторванные от своего
    приложения. Поэтому обходим тело документа напрямую.
    """
    import docx  # python-docx
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    document = docx.Document(str(path))
    body = document.element.body
    blocks: List[Block] = []
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            text = _normalize(Paragraph(child, document).text)
            if text:
                blocks.append(Block(text=text))
        elif tag == "tbl":
            for row in Table(child, document).rows:
                cells = [_normalize(cell.text) for cell in row.cells]
                line = " | ".join(cell for cell in cells if cell)
                if line:
                    blocks.append(Block(text=line))
    return blocks


def extract_plain(path: Path) -> List[Block]:
    text = _normalize(path.read_text(encoding="utf-8", errors="replace"))
    return [Block(text=line.strip()) for line in text.split("\n") if line.strip()]


def extract_blocks(path: Path) -> List[Block]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf(path)
    if suffix == ".docx":
        return extract_docx(path)
    if suffix in {".txt", ".md"}:
        return extract_plain(path)
    raise UnsupportedFormat(
        f"Формат {suffix or '(без расширения)'} не поддерживается. "
        f"Доступны: {', '.join(sorted(SUPPORTED_SUFFIXES))}"
    )


def chunk_blocks(
    blocks: Iterable[Block],
    chunk_size: int = 1200,
    overlap: int = 200,
    numbering: str = "auto",
) -> List[Chunk]:
    """Собрать абзацы во фрагменты ~chunk_size символов.

    Границы статей уважаются: новая «Статья N» всегда начинает новый фрагмент,
    чтобы ссылка на источник оставалась однозначной.
    """
    from .sections import LEVEL_CHAPTER, classify_heading  # circular-safe

    blocks = list(blocks)
    if numbering == "auto":
        numbering = detect_numbering(blocks)

    chunks: List[Chunk] = []
    buffer: List[str] = []
    buffer_len = 0
    current_locator = ""
    chunk_locator = ""
    chunk_page: Optional[int] = None
    chunk_section = ""
    chunk_chapter = ""

    def flush() -> None:
        nonlocal buffer, buffer_len
        text = "\n".join(buffer).strip()
        if text:
            chunks.append(
                Chunk(
                    text=text,
                    locator=chunk_locator,
                    page=chunk_page,
                    index=len(chunks),
                    section=chunk_section,
                    chapter=chunk_chapter,
                )
            )
        if overlap > 0 and text:
            # Хвост для перекрытия подрезается по границе абзаца, иначе цитата
            # в источнике начинается с обрывка слова.
            tail = text[-overlap:]
            newline = tail.find("\n")
            if newline != -1:
                tail = tail[newline + 1 :]
            buffer = [tail] if tail.strip() else []
            buffer_len = len(tail) if tail.strip() else 0
        else:
            buffer = []
            buffer_len = 0

    for block in blocks:
        if is_noise(block.text):
            continue

        # Заголовок части, раздела или главы закрывает фрагмент, но в текст не
        # попадает: он уже сохранён в метаданных, а отдельным фрагментом только
        # засоряет выдачу поиска.
        level = classify_heading(block.text)
        if level is not None and level <= LEVEL_CHAPTER:
            if buffer_len > 0:
                flush()
            buffer, buffer_len = [], 0
            heading_locator = detect_locator(block.text, numbering) or ""
            # Ссылаться на таблицу можно только через её приложение.
            current_locator = (
                heading_locator if heading_locator.lower().startswith("приложение") else ""
            )
            continue

        locator = detect_locator(block.text, numbering)
        is_boundary = bool(locator) and locator.lower().startswith(BOUNDARY_PREFIXES)
        if locator:
            current_locator = locator
        if is_boundary and buffer_len > 0:
            flush()
            buffer, buffer_len = [], 0  # без перекрытия через границу статьи
        if not buffer:
            chunk_locator = current_locator
            chunk_page = block.page
            chunk_section = block.section
            chunk_chapter = block.chapter
        buffer.append(block.text)
        buffer_len += len(block.text) + 1
        if buffer_len >= chunk_size:
            flush()
            chunk_locator = current_locator
            chunk_page = block.page
            chunk_section = block.section
            chunk_chapter = block.chapter

    if buffer and "".join(buffer).strip():
        text = "\n".join(buffer).strip()
        chunks.append(
            Chunk(
                text=text,
                locator=chunk_locator,
                page=chunk_page,
                index=len(chunks),
                section=chunk_section,
                chapter=chunk_chapter,
            )
        )

    return [c for c in chunks if len(c.text.strip()) >= 40]


def strip_preamble(blocks: List[Block]) -> List[Block]:
    """Отбросить титул и оглавление — всё до первой структурной единицы.

    Преамбула нормативного акта (название, реквизиты, текст приказа об
    утверждении) нормой не является, но короткие абзацы из неё стабильно
    всплывают в лексическом поиске.
    """
    from .sections import LEVEL_CHAPTER, classify_heading

    for position, block in enumerate(blocks):
        level = classify_heading(block.text)
        if level is not None and level <= LEVEL_CHAPTER:
            # Страховка от документов без привычной структуры.
            if position and position <= max(30, len(blocks) // 6):
                return blocks[position:]
            return blocks
    return blocks


def load_and_chunk(
    path: Path,
    chunk_size: int = 1200,
    overlap: int = 200,
    exclude_topics: Sequence[str] = (),
):
    """Прочитать файл, отбросить исключённые темы и нарезать на фрагменты.

    Возвращает (фрагменты, список выброшенных разделов).
    """
    from .sections import annotate_blocks, filter_blocks  # circular-safe

    blocks = annotate_blocks(strip_preamble(extract_blocks(path)))
    numbering = detect_numbering(blocks)
    kept, dropped = filter_blocks(blocks, exclude_topics)
    return (
        chunk_blocks(kept, chunk_size=chunk_size, overlap=overlap, numbering=numbering),
        dropped,
    )
